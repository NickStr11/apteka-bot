"""DOCX to plain text parser."""

from pathlib import Path
from typing import BinaryIO
import io

from docx import Document


def parse_docx(source: str | Path | BinaryIO) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        source: File path or file-like object
        
    Returns:
        Extracted plain text
    """
    try:
        doc = Document(source)
        text_parts: list[str] = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        return f"[DOCX parsing error: {e}]"


def parse_docx_bytes(content: bytes) -> str:
    """
    Extract text from DOCX bytes.
    
    Args:
        content: DOCX file content as bytes
        
    Returns:
        Extracted plain text
    """
    return parse_docx(io.BytesIO(content))
